"""
Genera la plantilla .docx GTH-F-56 de REGIS Colombia programáticamente.
Se ejecuta UNA VEZ para crear el template base.
Luego generate_report.py lo llena con datos.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(SCRIPT_DIR, "templates", "GTH-F-56_template.docx")


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shading_elem)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.makeelement(qn("w:tcBorders"), {})
    for edge, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            elem = borders.makeelement(qn(f"w:{edge}"), {
                qn("w:val"): "single",
                qn("w:sz"): val,
                qn("w:space"): "0",
                qn("w:color"): "000000",
            })
            borders.append(elem)
    tcPr.append(borders)


def set_cell_text(cell, text, bold=False, size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)


def create_template():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(9)

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # === HEADER TABLE ===
    header_table = doc.add_table(rows=3, cols=3)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    logo_cell = header_table.cell(0, 0)
    logo_cell.merge(header_table.cell(2, 0))
    set_cell_text(logo_cell, "REGIS COLOMBIA", bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    title_cell = header_table.cell(0, 1)
    title_cell.merge(header_table.cell(1, 1))
    set_cell_text(title_cell, "FORMATO SST: SEGUIMIENTO A RECOMENDACIONES\nMÉDICO LABORALES", bold=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    process_cell = header_table.cell(2, 1)
    set_cell_text(process_cell, "PROCESO: GESTIÓN DEL TALENTO HUMANO", bold=True, size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    set_cell_text(header_table.cell(0, 2), "Versión: 3.0", size=8)
    set_cell_text(header_table.cell(1, 2), "Fecha: 12/09/2019", size=8)
    set_cell_text(header_table.cell(2, 2), "Código: GTH-F-56", size=8)

    doc.add_paragraph()

    # === DATA FIELDS TABLE ===
    field_labels = [
        "FECHA",
        "NOMBRE DEL TRABAJADOR",
        "CÉDULA",
        "EPS",
        "CARGO",
        "EDAD",
        "PESO",
        "FECHA DE NACIMIENTO",
        "DEPENDENCIA",
        "EN TRATAMIENTO MÉDICO",
        "SEDE",
        "TIPO DE VINCULACIÓN LABORAL",
        "NOMBRE DEL JEFE INMEDIATO",
    ]

    placeholder_keys = [
        "{{fecha}}",
        "{{nombre_trabajador}}",
        "{{cedula}}",
        "{{eps}}",
        "{{cargo}}",
        "{{edad}}",
        "{{peso}}",
        "{{fecha_nacimiento}}",
        "{{dependencia}}",
        "{{en_tratamiento_medico}}",
        "{{sede}}",
        "{{tipo_vinculacion_laboral}}",
        "{{nombre_jefe_inmediato}}",
    ]

    data_table = doc.add_table(rows=len(field_labels), cols=2)
    data_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, placeholder) in enumerate(zip(field_labels, placeholder_keys)):
        set_cell_text(data_table.cell(i, 0), label, bold=True, size=9)
        set_cell_shading(data_table.cell(i, 0), "D9D9D9")
        set_cell_text(data_table.cell(i, 1), placeholder, size=9)

    # === TEXT BLOCKS ===
    text_sections = [
        ("Diagnóstico:", "{{diagnostico}}"),
        ("Tipo de tratamiento:", "{{tipo_tratamiento}}"),
        ("Funciones del cargo:", "{{funciones_cargo}}"),
    ]

    for title, placeholder in text_sections:
        doc.add_paragraph()
        tbl = doc.add_table(rows=2, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_cell_text(tbl.cell(0, 0), title, bold=True, size=9)
        set_cell_shading(tbl.cell(0, 0), "BFBFBF")
        set_cell_text(tbl.cell(1, 0), placeholder, size=9)

    # === PAGE 2: Seguimiento y compromisos ===
    doc.add_page_break()

    page2_sections = [
        ("Estado y cumplimiento de las recomendaciones médico laborales:", "{{recomendaciones_medico_laborales}}"),
        ("Observaciones:", "{{observaciones}}"),
        ("Compromiso de funcionario:", "{{compromiso_funcionario}}"),
        ("Compromiso de la entidad y áreas afines:", "{{compromiso_entidad}}"),
    ]

    for title, placeholder in page2_sections:
        tbl = doc.add_table(rows=2, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_cell_text(tbl.cell(0, 0), title, bold=True, size=9)
        set_cell_shading(tbl.cell(0, 0), "BFBFBF")
        set_cell_text(tbl.cell(1, 0), placeholder, size=9)
        doc.add_paragraph()

    # === PAGE 3: Firma ===
    doc.add_page_break()

    firma_tbl = doc.add_table(rows=2, cols=1)
    firma_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_text(firma_tbl.cell(0, 0), "Nombre, Cargo y Firma", bold=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(firma_tbl.cell(0, 0), "BFBFBF")
    set_cell_text(firma_tbl.cell(1, 0), "\n\n\n\n", size=9)

    os.makedirs(os.path.dirname(TEMPLATE_PATH), exist_ok=True)
    doc.save(TEMPLATE_PATH)
    print(f"Plantilla creada: {TEMPLATE_PATH}")
    return TEMPLATE_PATH


if __name__ == "__main__":
    create_template()
